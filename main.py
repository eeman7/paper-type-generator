import random as r
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from letters import letters

doc = open("C:/Users/adeto/Documents/STA203 PRACTICE QUESTIONS.docx", "rb")

docu = Document(doc)
raw_questions = []
for para in docu.paragraphs:
    raw_questions.append(para.text)

next_question = []
questions = []
for item in raw_questions:
    if item != "":
        next_question.append(item)
    else:
        questions.append(next_question)
        next_question = []

for i in range(4):
    document = Document()

    p1 = document.add_paragraph()
    p1_format = p1.paragraph_format
    p1_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    heading = p1.add_run(f"Paper Type {letters[i]}")
    font = heading.font
    font.size = Pt(20)
    heading.bold = True
    heading.underline = True

    r.shuffle(questions)

    for question in questions:
        p2 = document.add_paragraph(style='List Number')
        next_question = p2.add_run(f"{question[0]}")
        next_question.font.size = Pt(14)

        if len(question) > 1:
            for j in range(1, len(question)):
                p3 = document.add_paragraph(style='List Bullet')
                next_question = p3.add_run(f"{question[j]}")
                next_question.font.size = Pt(14)

    document.save(f"C:/Users/adeto/Documents/Paper Type {letters[i]}.docx")
